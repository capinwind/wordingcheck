import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
import requests
from io import BytesIO, StringIO

# Optional PDF parsing
try:
    import PyPDF2 as pypdf
except Exception:
    try:
        import pypdf
    except Exception:
        pypdf = None

# タイトル
st.title("品川学園PTA文章確認ミニツール")


st.markdown("---")
st.subheader("入力方法を選択")

# Quick open for rule document (Excel)
st.write("ルール文書（Excel）を開く")
rule_path = st.text_input("Enter Excel path or URL", key="rule_excel_input")
if st.button("Open rule document in excel"):
    rp = (rule_path or "").strip()
    if not rp:
        st.warning("パスまたはURLを入力してください")
    else:
        try:
            if rp.lower().startswith(("http://", "https://")):
                resp = requests.get(rp, timeout=15)
                resp.raise_for_status()
                df_rule = pd.read_excel(BytesIO(resp.content))
                st.success("ルール文書をURLから読み込みました")
                st.dataframe(df_rule.head())
            else:
                df_rule = pd.read_excel(rp)
                st.success("ルール文書をローカルから読み込みました")
                st.dataframe(df_rule.head())
        except Exception:
            st.error("cannot read the file, please check you link or the file behind your link")

# 選択肢（リンク / ローカル / テキスト貼り付け）
mode = st.selectbox(
    "入力方法を選んでください",
    ("1.リンクから開く (クラウド)", "2.ローカルから開く", "3.テキストを貼り付け")
)

if mode == "1.リンクから開く (クラウド)":
    st.write("クラウド上のファイルのURLを貼り付けてください。")
    url = st.text_input("ファイルのURLを貼り付け")
    if st.button("Open"):
        if url:
            # Try to read the URL into a DataFrame (CSV or Excel), or as text fallback.
            read_success = False
            # Try CSV
            try:
                df = pd.read_csv(url)
                st.success("ファイルを読み込みました (CSV)")
                st.dataframe(df.head())
                # store current data for analysis
                st.session_state["current_df"] = df
                st.session_state.pop("current_full_text", None)
                read_success = True
            except Exception:
                pass

            # Try Excel
            if not read_success:
                try:
                    df = pd.read_excel(url)
                    st.success("ファイルを読み込みました (Excel)")
                    st.dataframe(df.head())
                    st.session_state["current_df"] = df
                    st.session_state.pop("current_full_text", None)
                    read_success = True
                except Exception:
                    pass

            # Fetch raw bytes and try format-specific parsers (Excel, Word, PDF, text)
            if not read_success:
                try:
                    resp = requests.get(url, timeout=15)
                    resp.raise_for_status()
                    content = resp.content
                    ctype = (resp.headers.get("content-type") or "").lower()
                    lower = url.lower()

                    # Excel by content-type/extension or by attempting to read
                    try:
                        df = pd.read_excel(BytesIO(content))
                        st.success("ファイルを読み込みました (Excel)")
                        st.dataframe(df.head())
                        read_success = True
                    except Exception:
                        pass

                    # Word (.docx) - try to parse as docx
                    if not read_success:
                        try:
                            doc = Document(BytesIO(content))
                            text = "\n".join(p.text for p in doc.paragraphs)
                            if text.strip():
                                lines = text.splitlines()
                                df = pd.DataFrame({"text": lines})
                                df_full = pd.DataFrame({"full_text": [text]})
                                st.success("ファイルを読み込みました (Word .docx)")
                                st.dataframe(df.head())
                                st.dataframe(df_full)
                                # store for analysis
                                st.session_state["current_df"] = df
                                st.session_state["current_full_text"] = text
                                read_success = True
                        except Exception:
                            pass

                    # PDF - try parsing if pypdf is available
                    if not read_success:
                        try:
                            if pypdf is None:
                                st.warning("PDF読み取りには pypdf のインストールが必要です (pip install pypdf)")
                            else:
                                reader = pypdf.PdfReader(BytesIO(content))
                                pages = []
                                for page in reader.pages:
                                    try:
                                        page_text = page.extract_text()
                                    except Exception:
                                        page_text = None
                                    if page_text:
                                        pages.append(page_text)
                                if pages:
                                    text = "\n".join(pages)
                                    lines = text.splitlines()
                                    df = pd.DataFrame({"text": lines})
                                    df_full = pd.DataFrame({"full_text": [text]})
                                    st.success("ファイルを読み込みました (PDF)")
                                    st.dataframe(df.head())
                                    st.dataframe(df_full)
                                    st.session_state["current_df"] = df
                                    st.session_state["current_full_text"] = text
                                    read_success = True
                        except Exception:
                            pass

                    # Text fallback
                    if not read_success:
                        try:
                            try:
                                text = content.decode("utf-8")
                            except Exception:
                                text = content.decode("latin-1", errors="ignore")
                            lines = text.splitlines()
                            df = pd.DataFrame({"text": lines})
                            df_full = pd.DataFrame({"full_text": [text]})
                            st.success("テキストファイルを読み込みました")
                            st.dataframe(df.head())
                            st.dataframe(df_full)
                            st.session_state["current_df"] = df
                            st.session_state["current_full_text"] = text
                            read_success = True
                        except Exception:
                            read_success = False
                except Exception:
                    read_success = False

            if not read_success:
                st.error("cannot read the file, please check you link or the file behind your link")
        else:
            st.warning("URLを入力してください")
    


elif mode == "2.ローカルから開く":
    st.write("ローカルファイルを選択してください（プレビューを表示します）。")
    col1, col2 = st.columns([3, 1])
    # 表示用のファイルパステキストボックス（読み取り専用）
    file_path = st.session_state.get("local_file_path", "")
    col1.text_input("ファイルパス", value=file_path, key="local_file_path", disabled=True)
    if col2.button("Browse"):
        st.session_state["show_uploader"] = True

    if st.session_state.get("show_uploader", False):
        uploaded_file = st.file_uploader(
            "ファイルを選択してください",
            type=["txt", "csv", "xlsx", "xls", "docx"],
        )
        if uploaded_file is not None:
            st.session_state["local_file_path"] = uploaded_file.name
            st.success(f"選択されたファイル: {uploaded_file.name}")
            # 簡易プレビュー
            try:
                if uploaded_file.name.endswith(".csv"):
                    df = pd.read_csv(uploaded_file)
                    st.dataframe(df.head())
                    st.session_state["current_df"] = df
                    st.session_state.pop("current_full_text", None)
                elif uploaded_file.name.endswith((".xlsx", ".xls")):
                    df = pd.read_excel(uploaded_file)
                    st.dataframe(df.head())
                    st.session_state["current_df"] = df
                    st.session_state.pop("current_full_text", None)
                elif uploaded_file.name.endswith(".txt"):
                            text = uploaded_file.getvalue().decode("utf-8")
                            lines = text.splitlines()
                            df = pd.DataFrame({"text": lines})
                            df_full = pd.DataFrame({"full_text": [text]})
                            st.dataframe(df.head())
                            st.dataframe(df_full)
                            st.session_state["current_df"] = df
                            st.session_state["current_full_text"] = text
                elif uploaded_file.name.endswith(".docx"):
                            doc = Document(uploaded_file)
                            text = "\n".join(p.text for p in doc.paragraphs)
                            lines = text.splitlines()
                            df = pd.DataFrame({"text": lines})
                            df_full = pd.DataFrame({"full_text": [text]})
                            st.dataframe(df.head())
                            st.dataframe(df_full)
                            st.session_state["current_df"] = df
                            st.session_state["current_full_text"] = text
            except Exception as e:
                st.error(f"ファイルのプレビューに失敗しました: {e}")

else:  # テキストを貼り付け
    st.write("3.直接テキストを貼り付けてください。")
    pasted = st.text_area("テキストをここに貼り付け", height=200)
    if st.button("Use text"):
        if pasted.strip():
            # Put pasted text into both line-based DataFrame and full_text DataFrame
            lines = pasted.splitlines()
            df = pd.DataFrame({"text": lines})
            df_full = pd.DataFrame({"full_text": [pasted]})
            st.success("テキストを受け取りました。プレビュー：")
            st.dataframe(df.head())
            st.dataframe(df_full)
            st.session_state["current_df"] = df
            st.session_state["current_full_text"] = pasted
        else:
            st.warning("テキストを入力してください。")


# --- Analysis UI ---
st.markdown("---")
st.subheader("Analysis")
# Text area (~50 lines)
analysis_val = st.session_state.get("analysis_output", "")
# Read-only analysis output box (disabled=True); DO NOT bind to the same session_state key to allow updates
analysis_box = st.text_area("Analysis output", value=analysis_val, height=900, disabled=True)

from pathlib import Path

def load_df_word_rule():
    # Try session-state first
    df_wr = st.session_state.get("df_word_rule")
    if df_wr is not None:
        return df_wr
    # Try persisted file under Streamlit/data
    p = Path(__file__).resolve().parent / "data" / "df_word_rule.pkl"
    try:
        if p.exists():
            return pd.read_pickle(p)
    except Exception:
        return None
    return None

if st.button("analysis"):
    # Run analysis
    df_wr = load_df_word_rule()
    
    if df_wr is None:
        st.error("df_word_rule が見つかりません。'Wording Rules' ページでルールを読み込んでください。")
    else:
        # Ensure '誤表記' and '正表記' exist
        if "誤表記" not in df_wr.columns or "正表記" not in df_wr.columns:
            st.error("df_word_rule に '誤表記' または '正表記' 列が見つかりません")
        else:
            # Prepare target content string(s)
            current_full = st.session_state.get("current_full_text")
            current_df = st.session_state.get("current_df")
            # Create a searchable text: prefer full_text, else join all df cells
            if current_full:
                haystack = current_full
            elif current_df is not None:
                try:
                    haystack = "\n".join(current_df.astype(str).apply(lambda row: " ".join(row.values), axis=1).tolist())
                except Exception:
                    haystack = current_df.astype(str).to_string()
            else:
                st.warning("解析対象のデータがありません。ファイルを開くかテキストを貼り付けてください。")
                haystack = None

            if haystack is not None:
                out_lines = []
                for _, row in df_wr.iterrows():
                    typo = str(row.get("誤表記", ""))
                    corr = str(row.get("正表記", ""))
                    if not typo:
                        continue
                    try:
                        if typo in haystack:
                            out_lines.append(f"{typo} → {corr}")
                    except Exception:
                        # fallback to simple substring test on stringified haystack
                        if str(typo) in str(haystack):
                            out_lines.append(f"{typo} → {corr}")

                if out_lines:
                    prev = st.session_state.get("analysis_output", "")
                    new = prev + ("\n" if prev else "") + "\n".join(out_lines)
                    st.session_state["analysis_output"] = new
                    st.success("Analysis done — issues found and appended")
                else:
                    st.session_state["analysis_output"] = st.session_state.get("analysis_output", "")
                    st.success("Analysis done — no issues found")


