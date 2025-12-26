import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
import requests
from io import BytesIO, StringIO
# Optional PDF parsing
try:
    import PyPDF2 as pypdf
except Exception:
    try:
        import pypdf
    except Exception:
        pypdf = None

st.set_page_config(
        page_title="文章確認ツール",
)


st.title("品川学園PTA文章確認ミニツール")


# Quick open for rule document (Excel)
st.markdown("---")
st.subheader("1. 文章確認ルール")
# st.write(".   * 適用される文章ルールは別タブに表示しています。")
# st.write(".   * デフォルトは令和７年度版の文章ルールデータを使用しています。")
# st.write(".   * 別タブのテーブルデータを直接編集して、ルールの追加や編集できます")
# st.write(".   * 新しいルールファイルを読み込むことで、ルールデータを差し替えることもできます")

st.markdown("""
               :writing_hand: 適用される文章ルールは別タブに表示しています。</br>
               :writing_hand: デフォルトは令和７年度版の文章ルールデータを使用しています。</br>
               :writing_hand: 別タブのテーブルデータを直接編集して、ルールの追加や編集できます</br>
               :writing_hand: 新しいルールファイルを読み込むことで、ルールデータを差し替えることもできます</br>
            """, unsafe_allow_html=True)

st.markdown("---")
st.subheader("2. 文章確認対象")
st.markdown("""
               :green_heart: 想定される確認対象は、配信文章そのもの、又は添付ファイル。</br>
               :green_heart: 配信文章は、[3.テキスト貼りつけ] もしくは、WORDに入れて[1.ファイル(PC)]、両方OK。</br>
               :green_heart: ファイル(PC)は、ローカルPC保存のText/Word/Excelファイルなどを想定。</br>
               :broken_heart: ファイル(Link)は、Teamsにあるファイルの文章確認を想定だが、未対応。難しいかも。</br>
            """, unsafe_allow_html=True)


# st.write("**2-(1) 入力方法を選択してください。**")

# 選択肢（リンク / ローカル / テキスト貼り付け）
# mode = st.selectbox(
#     "入力方法を選んでください",
#     ("1.リンクから開く (クラウド)", "2.ローカルファイルを開く", "3.テキストを貼り付け"),
#     label_visibility="collapsed"
# )

page_names = ["1.ファイル(PC)",
              "2.ファイル(Link)", 
              "3.テキスト貼り付け"]
mode = st.radio(
    "**三つのモードから入力方法を選んでください**",
    page_names,
    # index=page_names.index(mode),
    # label_visibility="collapsed"
)

# st.write("    * 入力方法は三種類：①リンクから開く；②ローカルファイル　③テキストを直接貼り付け")
# st.write("    * 受付できるファイル形式：CSV, Excel, Word, PDF, テキストファイル")
# st.write("    ")

if mode == "1.ファイル(PC)":
    st.write("**PCローカルに保存されているファイルを選択してください**")
    # col1, col2 = st.columns([3, 1])
    # 表示用のファイルパステキストボックス（読み取り専用）
    file_path = st.session_state.get("local_file_path", "")
    # Show file path (disabled). Do NOT bind to session_state key to allow programmatic updates.
    # st.text_input("ファイルパス", value=file_path, disabled=True, label_visibility="collapsed")
    # if st.button("Click! (選択したファイルを読み込む)"):
    st.session_state["show_uploader"] = True
    # if st.session_state.get("show_uploader", False):
    uploaded_file = st.file_uploader(
        "ファイルを選択してください",
        type=["txt", "csv", "xlsx", "xls", "docx"],
        label_visibility="collapsed"
    )
    if uploaded_file is not None:
        st.session_state["local_file_path"] = uploaded_file.name
        st.success(f"選択されたファイル: {uploaded_file.name}")
        # 簡易プレビュー
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
                st.dataframe(df.head())
                st.session_state["current_df"] = df
                st.session_state.pop("current_full_text", None)
            elif uploaded_file.name.endswith((".xlsx", ".xls")):
                df = pd.read_excel(uploaded_file)
                st.dataframe(df.head())
                st.session_state["current_df"] = df
                st.session_state.pop("current_full_text", None)
            elif uploaded_file.name.endswith(".txt"):
                text = uploaded_file.getvalue().decode("utf-8")
                lines = text.splitlines()
                df = pd.DataFrame({"text": lines})
                df_full = pd.DataFrame({"full_text": [text]})
                # st.dataframe(df.head())
                st.dataframe(df_full,width=700,height="stretch",hide_index=True,row_height=100)
                st.session_state["current_df"] = df
                st.session_state["current_full_text"] = text
            elif uploaded_file.name.endswith(".docx"):
                doc = Document(uploaded_file)
                text = "\n".join(p.text for p in doc.paragraphs)
                lines = text.splitlines()
                df = pd.DataFrame({"text": lines})
                df_full = pd.DataFrame({"full_text": [text]})
                # st.dataframe(df.head())
                st.dataframe(df_full,width=700,hide_index=True,row_height=100)
                st.session_state["current_df"] = df
                st.session_state["current_full_text"] = text
        except Exception as e:
            st.error(f"ファイルのプレビューに失敗しました: {e}")

elif mode == "2.ファイル(Link)":
    st.write("**オンラインファイルのリンクを貼り付けて、読込みボタンを押してください。**")
    url = st.text_input("ファイルのURLを貼り付け",label_visibility="collapsed")
    if st.button("読込み"):
        if url:
            # Try to read the URL into a DataFrame (CSV or Excel), or as text fallback.
            read_success = False
            # Try CSV
            try:
                df = pd.read_csv(url)
                st.success("ファイルを読み込みました (CSV)")
                st.dataframe(df.head())
                # store current data for analysis
                st.session_state["current_df"] = df
                st.session_state.pop("current_full_text", None)
                read_success = True
            except Exception:
                pass

            # Try Excel
            if not read_success:
                try:
                    df = pd.read_excel(url)
                    st.success("ファイルを読み込みました (Excel)")
                    st.dataframe(df.head())
                    st.session_state["current_df"] = df
                    st.session_state.pop("current_full_text", None)
                    read_success = True
                except Exception:
                    pass

            # Fetch raw bytes and try format-specific parsers (Excel, Word, PDF, text)
            if not read_success:
                try:
                    resp = requests.get(url, timeout=15)
                    resp.raise_for_status()
                    content = resp.content
                    ctype = (resp.headers.get("content-type") or "").lower()
                    lower = url.lower()

                    # Excel by content-type/extension or by attempting to read
                    try:
                        df = pd.read_excel(BytesIO(content))
                        st.success("ファイルを読み込みました (Excel)")
                        st.dataframe(df.head())
                        read_success = True
                    except Exception:
                        pass

                    # Word (.docx) - try to parse as docx
                    if not read_success:
                        try:
                            doc = Document(BytesIO(content))
                            text = "\n".join(p.text for p in doc.paragraphs)
                            if text.strip():
                                lines = text.splitlines()
                                df = pd.DataFrame({"text": lines})
                                df_full = pd.DataFrame({"full_text": [text]})
                                st.success("ファイルを読み込みました (Word .docx)")
                                # st.dataframe(df.head())
                                st.dataframe(df_full,width=700,hide_index=True,row_height=100)
                                # store for analysis
                                st.session_state["current_df"] = df
                                st.session_state["current_full_text"] = text
                                read_success = True
                        except Exception:
                            pass

                    # PDF - try parsing if pypdf is available
                    if not read_success:
                        try:
                            if pypdf is None:
                                st.warning("PDF読み取りには pypdf のインストールが必要です (pip install pypdf)")
                            else:
                                reader = pypdf.PdfReader(BytesIO(content))
                                pages = []
                                for page in reader.pages:
                                    try:
                                        page_text = page.extract_text()
                                    except Exception:
                                        page_text = None
                                    if page_text:
                                        pages.append(page_text)
                                if pages:
                                    text = "\n".join(pages)
                                    lines = text.splitlines()
                                    df = pd.DataFrame({"text": lines})
                                    df_full = pd.DataFrame({"full_text": [text]})
                                    st.success("ファイルを読み込みました (PDF)")
                                    # st.dataframe(df.head())
                                    st.dataframe(df_full,width=700,hide_index=True,row_height=100)
                                    st.session_state["current_df"] = df
                                    st.session_state["current_full_text"] = text
                                    read_success = True
                        except Exception:
                            pass

                    # Text fallback
                    if not read_success:
                        try:
                            try:
                                text = content.decode("utf-8")
                            except Exception:
                                text = content.decode("latin-1", errors="ignore")
                            lines = text.splitlines()
                            df = pd.DataFrame({"text": lines})
                            df_full = pd.DataFrame({"full_text": [text]})
                            st.success("テキストファイルを読み込みました")
                            # st.dataframe(df.head())
                            st.dataframe(df_full,width=700,hide_index=True,row_height=100)
                            st.session_state["current_df"] = df
                            st.session_state["current_full_text"] = text
                            read_success = True
                        except Exception:
                            read_success = False
                except Exception:
                    read_success = False

            if not read_success:
                st.error("cannot read the file, please check you link or the file behind your link")
        else:
            st.warning("URLを入力してください")
    
else:  # テキストを貼り付け
    st.write("**下のテキストボックスに文章を入力/貼り付けて、読込みボタンを押してください。**")
    pasted = st.text_area("テキストをここに貼り付け", height=150, label_visibility="collapsed")
    if st.button("読込み"):
        if pasted.strip():
            # Put pasted text into both line-based DataFrame and full_text DataFrame
            lines = pasted.splitlines()
            df = pd.DataFrame({"text": lines})
            df_full = pd.DataFrame({"full_text": [pasted]})
            st.success("テキストを読み込めました。プレビュー：")
            # st.dataframe(df.head())
            st.dataframe(df_full,width=700,hide_index=True,row_height=100)
            st.session_state["current_df"] = df
            st.session_state["current_full_text"] = pasted
        else:
            st.warning("テキストを入力してください。")


# --- Analysis UI ---
st.markdown("---")
st.subheader("3. 文章確認実行")
# button_analysis = st.button("Just do it!!（改善する前は2回クリックしてください）")
# Initialize analysis output once immediately when button is pressed so the text area reflects it on this run
# if button_analysis:
#     st.session_state.setdefault("analysis_output", "")
# Text area will be shown below after analysis runs so updates appear in the same click

from pathlib import Path

def load_df_word_rule():
    # Try session-state first
    df_wr = st.session_state.get("df_word_rule")
    if df_wr is not None:
        return df_wr
    # Try persisted file under Streamlit/data
    p = Path(__file__).resolve().parent / "data" / "df_word_rule.pkl"
    try:
        if p.exists():
            return pd.read_pickle(p)
    except Exception:
        return None
    return None

if st.button("16355!!"):
    # Clear previous analysis output so the box shows only current results
    st.session_state["analysis_output"] = ""
    # st.session_state.setdefault("analysis_output", "")
    # Run analysis
    df_wr = load_df_word_rule()
    
    if df_wr is None:
        st.error("df_word_rule が見つかりません。'Wording Rules' ページでルールを読み込んでください。")
    else:
        # Ensure '誤表記' and '正表記' exist
        if "誤表記" not in df_wr.columns or "正表記" not in df_wr.columns:
            st.error("df_word_rule に '誤表記' または '正表記' 列が見つかりません")
        else:
            # Prepare target content string(s)
            current_full = st.session_state.get("current_full_text")
            current_df = st.session_state.get("current_df")
            # Create a searchable text: prefer full_text, else join all df cells
            if current_full:
                haystack = current_full
            elif current_df is not None:
                try:
                    haystack = "\n".join(current_df.astype(str).apply(lambda row: " ".join(row.values), axis=1).tolist())
                except Exception:
                    haystack = current_df.astype(str).to_string()
            else:
                st.warning("解析対象のデータがありません。ファイルを開くかテキストを貼り付けてください。")
                haystack = None

            if haystack is not None:
                out_lines = []
                for _, row in df_wr.iterrows():
                    typo = str(row.get("誤表記", ""))
                    corr = str(row.get("正表記", ""))
                    if not typo:
                        continue
                    try:
                        if typo in haystack:
                            out_lines.append(f"{typo} → {corr}")
                    except Exception:
                        # fallback to simple substring test on stringified haystack
                        if str(typo) in str(haystack):
                            out_lines.append(f"{typo} → {corr}")

                if out_lines:
                    prev = st.session_state.get("analysis_output", "")
                    new = prev + ("\n" if prev else "") + "\n".join(out_lines)
                    st.session_state["analysis_output"] = new
                    st.success("Analysis done — issues found and appended")
                else:
                    st.session_state["analysis_output"] = st.session_state.get("analysis_output", "")
                    st.success("Analysis done — no issues found")

# Display analysis output after processing so the text area reflects changes immediately
analysis_val = st.session_state.get("analysis_output", "")
st.text_area("Analysis output", value=analysis_val, height=500, disabled=True)